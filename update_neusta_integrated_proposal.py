from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path("/Users/johansandersnas/Downloads/NORC_Team_Neusta_90_90_Retainer_Proposal_SCALE_85000_final.docx")
OUTPUT = Path("/Users/johansandersnas/Downloads/NORC_Team_Neusta_Integrated_90_90_Proposal_updated.docx")

YELLOW = "FFCD05"
DARK = "353031"
MUTED = "6F6B6D"
LIGHT_YELLOW = "FFF4C2"
RULE = "DAD7D2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=RULE, val="single", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), val)
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=None, bold=None, color=None, all_caps=False):
    run.font.name = "Avenir Next"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Avenir Next")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if all_caps:
        run.font.all_caps = True


def add_p(doc, text="", style="Normal", size=None, bold=None, color=DARK, after=5):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_bullets(doc, items):
    for item in items:
        add_p(doc, f"•  {item}", size=9.6, after=3)


def add_callout(doc, title, body, fill=LIGHT_YELLOW):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table, 8800)
    set_cell_width(table.cell(0, 0), 120)
    set_cell_width(table.cell(0, 1), 8680)
    set_cell_shading(table.cell(0, 0), YELLOW)
    set_cell_shading(table.cell(0, 1), fill)
    for cell in table.row_cells(0):
        set_cell_border(cell, color=fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = table.cell(0, 1).paragraphs[0]
    r = p.add_run(title)
    set_run(r, size=9, bold=True, color=DARK)
    p2 = table.cell(0, 1).add_paragraph()
    r2 = p2.add_run(body)
    set_run(r2, size=9, color=DARK)
    add_p(doc, "", after=6)


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table, sum(widths))
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_width(cell, widths[i])
        set_cell_shading(cell, "F4F1EA")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run(r, size=8.5, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_cell_width(cell, widths[i])
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run(r, size=8.5, color=DARK)
    add_p(doc, "", after=8)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build():
    doc = Document(SOURCE)
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    doc.core_properties.title = "Team Neusta and NORC Integrated 90+90 Proposal"
    doc.core_properties.subject = "Combined Scenario A and B collaboration roadmap"
    doc.core_properties.author = "NORC"

    add_p(doc, "PROPOSAL", style="SmallCaps", size=7, bold=True, color=MUTED, after=8)
    add_p(doc, "Team Neusta x NORC / Viable", size=22, bold=True, after=4)
    add_p(doc, "Integrated 90+90 Market Collaboration", size=17, bold=True, after=4)
    add_p(
        doc,
        "Scenario B as the operational starting point. Scenario A as the strategic joint set-up.",
        size=10.5,
        color=MUTED,
        after=12,
    )
    add_callout(
        doc,
        "Core clarification",
        "This proposal does not treat Scenario A and Scenario B as two separate alternatives. It integrates both into one coherent 90+90 collaboration: Team Neusta as NORC's client for Nordic market entry, and a jointly governed set-up where Team Neusta, NORC and Viable develop a shared toolkit, collaboration model and future market opportunities together.",
    )
    add_simple_table(
        doc,
        ["PREPARED FOR", "PREPARED BY", "STATUS"],
        [["Team Neusta", "NORC - A Viable Solution", "Updated combined proposal for discussion"]],
        [2500, 3000, 3300],
    )

    add_p(doc, "1. How We See Scenario A and Scenario B", style="Heading 1")
    add_p(
        doc,
        "André's question is important because it clarifies the strategic nature of the collaboration. NORC's view is that the two scenarios should not be handled as competing tracks. They describe two layers of the same engagement.",
    )
    add_simple_table(
        doc,
        ["Layer", "Meaning", "Role in the 90+90 collaboration"],
        [
            [
                "Scenario B",
                "Team Neusta as NORC's client, with the focus on building Team Neusta's presence in Sweden and the Nordics.",
                "This is the operational starting point for the first 90 days: clear scope, delivery rhythm, outputs and commercial responsibility.",
            ],
            [
                "Scenario A",
                "A joint set-up where both sides contribute their respective strengths and develop a shared toolkit and collaboration model together.",
                "This is the strategic frame that shapes the whole 90+90 period, including the joint toolbox, governance, MoU and future market opportunities.",
            ],
        ],
        [1900, 3300, 3600],
    )
    add_p(
        doc,
        "In practical terms, Team Neusta can engage NORC through a client/retainer model while the work itself is designed to build a joint set-up. The client track creates delivery clarity; the joint set-up creates strategic direction.",
    )

    add_p(doc, "2. Strategic Intent", style="Heading 1")
    add_p(
        doc,
        "The ambition is to use the first 90 days to create the Neusta-NORC Opportunity Space and the first version of a joint market entry toolbox. The second 90 days should then move the collaboration into controlled activation, market testing and joint development of the broader Team Neusta / NORC / Viable platform.",
    )
    add_bullets(
        doc,
        [
            "Team Neusta contributes digital transformation capacity, engineering depth, delivery scale and German market strength.",
            "NORC contributes Nordic market understanding, positioning, communication, trust-building and strategic activation capability.",
            "Viable contributes resilience thinking, operational structure and the broader strategic frame for long-term cooperation.",
            "Together, the parties develop a shared way to translate complex digital capability into market relevance in the Nordics and, in the next step, into reciprocal opportunities in Germany.",
        ],
    )

    add_p(doc, "3. Integrated 90+90 Roadmap", style="Heading 1")
    add_simple_table(
        doc,
        ["Period", "Primary focus", "Main outputs"],
        [
            [
                "First 90 days",
                "Nordic market entry foundation for Team Neusta, structured as a SCALE Capability as a Service engagement.",
                "Opportunity Space Map, Nordic positioning hypothesis, joint key offer, audience and sector map, messaging architecture, short presentation material, conversation guide, activation backlog and Phase 2 plan.",
            ],
            [
                "Second 90 days",
                "Controlled market activation and strategic refinement, including the joint set-up dimension and the NORC / Viable with Team Neusta in Germany perspective.",
                "Test dialogues, selected prospect and partner mapping, sector-specific offer packaging, thought-leadership concepts, joint governance model, MoU finalisation and longer-term go-to-market model.",
            ],
        ],
        [1700, 3550, 3550],
    )
    add_callout(
        doc,
        "Suggested June milestone",
        "NORC proposes a kick-off in Bremen in June to align the parties, agree the 90+90 roadmap, define the shared governance model and prepare a non-binding Memorandum of Understanding for internal Board-level anchoring.",
    )

    add_page_break(doc)
    add_p(doc, "4. Phase 1 - First 90 Days", style="Heading 1")
    add_p(doc, "NEUSTA-NORC OPPORTUNITY SPACE & TOOLBOX FOUNDATION", style="SmallCaps", size=7.5, bold=True, color=MUTED)
    add_p(
        doc,
        "Purpose: to create a clear, visual and actionable view of the strongest Team Neusta opportunities in Sweden and the Nordics, while developing the first version of the shared toolbox that can later support the joint set-up.",
    )
    add_bullets(
        doc,
        [
            "Discovery and strategic framing of Team Neusta's Nordic ambition, sectors, proof points, buyer groups and available case material.",
            "Opportunity Space Map showing strongest entry points, barriers, audiences, relevance logic and first moves.",
            "Nordic positioning hypothesis and message architecture.",
            "Joint key offer and short presentation material for meetings and internal alignment.",
            "Activation backlog and recommendation for the second 90 days.",
        ],
    )

    add_p(doc, "5. Phase 2 - Next 90 Days", style="Heading 1")
    add_p(doc, "MARKET ACTIVATION, STRATEGIC REFINEMENT AND JOINT SET-UP", style="SmallCaps", size=7.5, bold=True, color=MUTED)
    add_p(
        doc,
        "Phase 2 should only be activated once the parties have agreed on the opportunity space, positioning and toolbox from Phase 1. The purpose is to move from foundation to controlled activation while also developing the joint set-up and reciprocal Germany/Nordic collaboration logic.",
    )
    add_bullets(
        doc,
        [
            "selected prospect, partner and ecosystem mapping;",
            "relationship and stakeholder mapping in Sweden, the Nordics and relevant German contexts;",
            "sector-specific offer packaging and content development;",
            "support in first customer, partner or ecosystem conversations;",
            "refinement of the toolbox based on live market feedback;",
            "development of a joint governance model and MoU structure;",
            "preparation of a longer-term Nordic/German collaboration model.",
        ],
    )

    add_p(doc, "6. Working Rhythm and Governance", style="Heading 1")
    add_p(
        doc,
        "The work will be managed as a senior group delivery rather than as a staffing model. Team Neusta buys access to an embedded NORC capability with the right combination of strategy, positioning, communication, facilitation and selected specialist input.",
    )
    add_bullets(
        doc,
        [
            "weekly working session between NORC and Team Neusta representatives;",
            "shared backlog of priorities, decisions and deliverables;",
            "two strategic alignment sessions per month;",
            "end-of-phase review after the first 90 days;",
            "joint steering discussion before activating the second 90 days;",
            "MoU preparation as a non-binding strategic intent document for respective Boards.",
        ],
    )

    add_p(doc, "7. Commercial Frame", style="Heading 1")
    add_p(
        doc,
        "NORC proposes that the combined 90+90 collaboration starts at SCALE. This reflects that Team Neusta's Nordic ambition is not an early start-up question, but a structured effort to scale an established digital transformation group into a new geographical, linguistic and cultural market.",
    )
    add_simple_table(
        doc,
        ["Element", "Commercial logic"],
        [
            ["Monthly SCALE partner price", "85,000 SEK per month excl. VAT."],
            ["First 90 days", "255,000 SEK excl. VAT."],
            ["Full 90+90 planning frame", "510,000 SEK excl. VAT, provided the second 90-day scope remains within the logic described in this proposal."],
            ["Continuation", "The second 90 days should be confirmed after the Phase 1 review, with scope and delivery emphasis agreed before continuation."],
            ["Travel and third-party costs", "Travel to Bremen or other locations, third-party production and paid media are excluded unless explicitly agreed."],
        ],
        [2800, 6000],
    )
    add_p(doc, "Included", style="Heading 2")
    add_bullets(
        doc,
        [
            "strategic advisory and senior facilitation;",
            "Nordic market entry positioning;",
            "narrative and message development;",
            "Opportunity Space and joint toolbox development;",
            "selected prospect, sector and ecosystem mapping;",
            "MoU preparation and joint governance framing;",
            "coordination, prioritisation and end-of-phase recommendations.",
        ],
    )
    add_p(doc, "Excluded and quoted separately", style="Heading 2")
    add_bullets(
        doc,
        [
            "major campaign production, website development, paid media and PR execution;",
            "film production, advanced design systems and technical implementation;",
            "travel costs, third-party costs and local staffing beyond the agreed retainer scope.",
        ],
    )

    add_page_break(doc)
    add_p(doc, "Appendix - Memorandum of Understanding Outline", style="Heading 1")
    add_p(
        doc,
        "NORC can prepare a non-binding Memorandum of Understanding for Team Neusta, NORC and Viable to share with their respective Boards. The purpose is not to create a legal partnership agreement at this stage, but to clarify strategic intent, collaboration principles and governance before or during the June kick-off in Bremen.",
    )
    add_simple_table(
        doc,
        ["MoU area", "Suggested content"],
        [
            ["Purpose", "Clarify the ambition to develop a joint set-up, shared toolkit and long-term collaboration model."],
            ["Parties", "Team Neusta, NORC and Viable, with roles and contribution areas described at a strategic level."],
            ["Scope", "Nordic market entry for Team Neusta, joint toolbox development, and exploration of reciprocal NORC / Viable opportunities with Team Neusta in Germany."],
            ["Governance", "Named senior owners, meeting cadence, decision points and Board-level anchoring."],
            ["Outputs", "Opportunity Space, joint key offer, market entry toolbox, activation roadmap and longer-term collaboration model."],
            ["Commercial principle", "The 90+90 SCALE retainer funds the agreed working scope; any separate ventures, production or implementation work are agreed separately."],
            ["Nature", "Non-binding strategic intent document, unless the parties later decide to convert parts into a formal agreement."],
        ],
        [2300, 6500],
    )

    add_p(doc, "Proposed Next Step", style="Heading 1")
    add_p(
        doc,
        "NORC suggests a short call with André Conin, Claudia Krüger, Johan Sandersnäs and Jonathan Converse to confirm the combined A/B interpretation, agree the June Bremen kick-off ambition and align on whether the first 90 days should begin as the operational start of the integrated 90+90 roadmap.",
    )
    add_callout(
        doc,
        "Closing statement",
        "Team Neusta has the digital transformation capacity, engineering depth and delivery scale required to become a relevant partner in the Nordic market. NORC can help make that capacity locally meaningful, trusted and commercially actionable. The proposed 90+90 model gives the collaboration a practical starting point and a shared strategic direction.",
        fill="F7F5EF",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
