#!/usr/bin/env python3
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
CONTENT = ROOT / "content"
ASSETS = ROOT / "assets"

NORC_YELLOW = "FFCD05"
NORC_DARK = "353031"
NORC_MUTED = "6F6B6D"
NORC_RULE = "DAD7D2"
FONT = "Avenir Next"
CONTENT_WIDTH_DXA = 9240


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key, value in kwargs[edge].items():
                element.set(qn("w:{}".format(key)), str(value))


def set_table_width(table, width_dxa):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_paragraph_bottom_border(paragraph, color=NORC_RULE, size="6", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_run_font(run, size=None, bold=False, color=NORC_DARK, all_caps=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if all_caps:
        run.font.all_caps = True


def add_text(paragraph, text, size=10.5, bold=False, color=NORC_DARK):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NORC_DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, before, after in [
        ("Title", 21, 0, 10),
        ("Heading 1", 12.8, 14, 6),
        ("Heading 2", 10.8, 10, 4),
        ("Heading 3", 10.2, 8, 3),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NORC_DARK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc: Document):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, "NORC / A VIABLE SOLUTION", size=7.5, color=NORC_DARK)
        p.paragraph_format.space_before = Pt(4)


def add_brand_header(doc: Document, data: dict):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table, CONTENT_WIDTH_DXA)
    widths = [980, 3050, CONTENT_WIDTH_DXA - 980 - 3050]
    for i, width in enumerate(widths):
        set_cell_width(table.cell(0, i), width)
        table.cell(0, i).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        for side in ("top", "left", "bottom", "right"):
            set_cell_border(table.cell(0, i), **{side: {"val": "nil"}})

    logo_cell = table.cell(0, 0)
    p = logo_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(ASSETS / "norc-logo.png"), width=Inches(0.66))

    brand = table.cell(0, 1).paragraphs[0]
    brand.paragraph_format.space_before = Pt(7)
    brand.paragraph_format.line_spacing = 1.0
    add_text(brand, "Nordic\nResilience\nCommunication", size=12.5, color=NORC_DARK)

    doc_type = table.cell(0, 2).paragraphs[0]
    doc_type.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc_type.paragraph_format.space_before = Pt(16)
    add_text(doc_type, data["document_type"], size=11, color=NORC_DARK)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(18)
    set_paragraph_bottom_border(spacer, size="4", space="1")


def add_metadata(doc: Document, data: dict):
    table = doc.add_table(rows=3, cols=4)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    widths = [1340, 2300, 3720, CONTENT_WIDTH_DXA - 1340 - 2300 - 3720]
    set_table_width(table, sum(widths))
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": NORC_RULE},
                bottom={"val": "single", "sz": "4", "color": NORC_RULE},
                left={"val": "nil"},
                right={"val": "nil"},
            )

    labels = ["Date", "Client", "Client contact", "Project number"]
    values = [data["date"], data["client"], data["client_contact"], data["project_number"]]
    for i, label in enumerate(labels):
        p = table.cell(0, i).paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        add_text(p, label, size=9.2, bold=True)
        p2 = table.cell(1, i).paragraphs[0]
        p2.paragraph_format.space_after = Pt(2)
        add_text(p2, values[i], size=9.2)

    merged = table.cell(2, 0).merge(table.cell(2, 3))
    p = merged.paragraphs[0]
    add_text(p, "Project", size=9.2, bold=True)
    p2 = merged.add_paragraph()
    add_text(p2, data["project"], size=9.2)


def parse_markdown(path: Path):
    lines = path.read_text(encoding="utf-8").splitlines()
    result = []
    for raw in lines:
        line = raw.strip()
        if not line:
            result.append(("blank", ""))
        elif line.startswith("## "):
            result.append(("h2", line[3:]))
        elif line.startswith("# "):
            result.append(("h1", line[2:]))
        else:
            result.append(("p", line))
    return result


def add_section_from_markdown(doc: Document, module_name: str):
    path = CONTENT / "{}.md".format(module_name)
    if not path.exists():
        raise FileNotFoundError("Missing module: {}".format(path))

    previous_blank = False
    for kind, text in parse_markdown(path):
        if kind == "blank":
            previous_blank = True
            continue
        if kind == "h1":
            p = doc.add_paragraph(style="Heading 1")
            if "Phase " in text:
                p.paragraph_format.keep_with_next = True
            add_text(p, text, size=11.5, bold=True)
        elif kind == "h2":
            p = doc.add_paragraph(style="Heading 2")
            add_text(p, text, size=10.2, bold=True)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4 if previous_blank else 0)
            p.paragraph_format.space_after = Pt(5)
            if text.endswith(":") or text in {"Investment (indicative)", "Optional success-based component"}:
                add_text(p, text, size=10.0, bold=text.endswith(":"))
            elif text[:2].isdigit() or text[:1].isdigit():
                add_text(p, text, size=10.0, bold=True)
            else:
                add_text(p, text, size=10.0)
        previous_blank = False


def add_intro(doc: Document, data: dict):
    p = doc.add_paragraph()
    add_text(p, data["headline"], size=10.3, bold=True)
    set_paragraph_bottom_border(p, size="4", space="2")
    for line in data.get("intro", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        add_text(p, line, size=10.6)


def build(project_path: Path, output_path: Path):
    data = json.loads(project_path.read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc)
    add_brand_header(doc, data)
    add_metadata(doc, data)
    add_intro(doc, data)
    for module in data["modules"]:
        add_section_from_markdown(doc, module)
    add_footer(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main(argv):
    if len(argv) != 3:
        print("Usage: build_offer.py <project.json> <output.docx>", file=sys.stderr)
        return 2
    build(Path(argv[1]).resolve(), Path(argv[2]).resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
